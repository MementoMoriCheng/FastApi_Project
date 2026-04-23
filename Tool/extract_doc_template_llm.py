import json
import os
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import Counter, defaultdict

import httpx
from docx import Document
from openai import OpenAI


DEFAULT_LLM_MODEL = os.getenv("DOC_TEMPLATE_LLM_MODEL", "qwen-turbo").strip()
DEFAULT_LLM_BASE_URL = os.getenv(
    "DOC_TEMPLATE_LLM_BASE_URL",
    "https://dashscope.aliyuncs.com/compatible-mode/v1",
).strip()
DEFAULT_LLM_API_KEY = os.getenv("DOC_TEMPLATE_LLM_API_KEY", "sk-0b4c79748dc04bd8bebbadde4a566fba").strip()
DEFAULT_LLM_TIMEOUT = float(os.getenv("DOC_TEMPLATE_LLM_TIMEOUT", "90"))
DEFAULT_LLM_CONNECT_TIMEOUT = float(os.getenv("DOC_TEMPLATE_LLM_CONNECT_TIMEOUT", "10"))
DEFAULT_LLM_MAX_RETRIES = int(os.getenv("DOC_TEMPLATE_LLM_MAX_RETRIES", "0"))
DEFAULT_BLOCK_SIZE = int(os.getenv("DOC_TEMPLATE_LLM_BLOCK_SIZE", "3"))
DEFAULT_BLOCK_STEP = int(os.getenv("DOC_TEMPLATE_LLM_BLOCK_STEP", "3"))
DEFAULT_LLM_MAX_WORKERS = int(os.getenv("DOC_TEMPLATE_LLM_MAX_WORKERS", "4"))

ALLOWED_CLASSIFICATIONS = {"template_fixed", "user_filled", "blank_slot"}
REPLACEABLE_CLASSIFICATIONS = {"user_filled", "blank_slot"}


def _clean_text(text):
    return str(text or "").replace("\xa0", " ").strip()


def _get_vmerge_state(cell):
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    vmerge = tc_pr.vMerge
    if vmerge is None:
        return None
    return vmerge.val or "continue"


def _compute_span_height(table, row_index, grid_index, span_width):
    current_row = table.rows[row_index]
    current_cell = current_row.cells[grid_index]
    current_vmerge = _get_vmerge_state(current_cell)
    if current_vmerge != "restart":
        return 1

    span_height = 1
    for next_row_index in range(row_index + 1, len(table.rows)):
        next_row = table.rows[next_row_index]
        if grid_index >= len(next_row.cells):
            break
        next_cell = next_row.cells[grid_index]
        next_vmerge = _get_vmerge_state(next_cell)
        if next_vmerge != "continue":
            break
        if span_width > 1:
            end_index = min(grid_index + span_width, len(next_row.cells))
            if not all(
                _get_vmerge_state(next_row.cells[i]) in {"restart", "continue"}
                for i in range(grid_index, end_index)
            ):
                break
        span_height += 1
    return span_height


def _iter_row_cells_with_spans(table, row_index):
    row = table.rows[row_index]
    raw_cells = list(row.cells)
    grid_index = 0
    seen = set()
    while grid_index < len(raw_cells):
        cell = raw_cells[grid_index]
        tc_id = id(cell._tc)
        if tc_id in seen:
            grid_index += 1
            continue
        seen.add(tc_id)
        next_index = grid_index + 1
        while next_index < len(raw_cells) and id(raw_cells[next_index]._tc) == tc_id:
            next_index += 1
        span_width = next_index - grid_index
        span_height = _compute_span_height(table, row_index, grid_index, span_width)
        yield grid_index, cell, span_width, span_height
        grid_index = next_index


def _parse_docx_tables(docx_path):
    doc = Document(docx_path)
    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row_index, _row in enumerate(table.rows):
            cells = []
            for cell_index, cell, span_width, span_height in _iter_row_cells_with_spans(table, row_index):
                text = _clean_text("\n".join(paragraph.text for paragraph in cell.paragraphs))
                cells.append(
                    {
                        "cell_index": cell_index,
                        "text": text,
                        "span_width": span_width,
                        "span_height": span_height,
                    }
                )
            rows.append({"row_index": row_index, "cells": cells})
        tables.append({"table_index": table_index, "rows": rows})
    return tables


def _parse_docx_paragraphs(docx_path):
    doc = Document(docx_path)
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = _clean_text(paragraph.text)
        if not text:
            continue
        paragraphs.append({"index": index, "text": text})
    return paragraphs


def _chunk_table_rows(rows, block_size=DEFAULT_BLOCK_SIZE, block_step=DEFAULT_BLOCK_STEP):
    if not rows:
        return []
    block_size = max(2, int(block_size or DEFAULT_BLOCK_SIZE))
    block_step = max(1, int(block_step or DEFAULT_BLOCK_STEP))
    blocks = []
    start = 0
    block_index = 0
    while start < len(rows):
        chunk = rows[start : start + block_size]
        if not chunk:
            break
        blocks.append({"block_index": block_index, "rows": chunk})
        if len(chunk) < block_size:
            break
        start += block_step
        block_index += 1
    return blocks


def _build_block_payload(table_index, block_index, rows):
    return {
        "task": "classify_table_block",
        "table_index": table_index,
        "block_index": block_index,
        "rows": rows,
    }


def _build_block_messages(block_payload):
    return [
        {
            "role": "system",
            "content": (
                "你是一个 DOCX 表格模板识别助手。"
                "你会看到同一个表格中的连续几行。"
                "请判断每个单元格属于以下三类之一："
                "template_fixed 表示模板原本就有的固定文本、表头、标签、说明项、枚举标题；"
                "user_filled 表示用户已经填写进去的值、数据、名称、日期、金额、编号、枚举结果；"
                "blank_slot 表示模板中待填写的空白单元格或空白值区域。"
                "只返回 JSON，不要输出 markdown，不要添加额外解释。"
            ),
        },
        {
            "role": "user",
            "content": json.dumps(
                {
                    "instructions": [
                        "结合同一个 block 中多行之间的关系判断，不要只看单行。",
                        "优先识别列头、标签、空白模板格、以及已经填写的数据值。",
                        "如果一行中出现连续的短数字如 1 到 12，且下一行是对应的数据值，这些短数字通常是列头，应该判为 template_fixed，不是 blank_slot。",
                        "必须为每个输入 cell_index 返回一个 classification。",
                        "classification 只能是 template_fixed、user_filled、blank_slot 之一。",
                        "输出格式必须为 {rows: [{row_index, cells: [{cell_index, classification, reason}]}], notes: ''}。",
                    ],
                    "block": block_payload,
                },
                ensure_ascii=False,
            ),
        },
    ]


def _build_paragraph_messages(paragraph_payload):
    return [
        {
            "role": "system",
            "content": (
                "你是一个 DOCX 模板识别助手。"
                "请判断段落属于以下三类之一："
                "template_fixed 表示模板原本就有的固定说明、标题、标签、提示文本；"
                "user_filled 表示用户已经填写进去的具体内容；"
                "blank_slot 表示虽然有占位形式，但本质上是待填写区域。"
                "只返回 JSON，不要输出 markdown。"
            ),
        },
        {
            "role": "user",
            "content": json.dumps(
                {
                    "instructions": [
                        "如果段落是固定说明、标题、或模板原文，判为 template_fixed。",
                        "如果段落是已填写的具体描述、名称、地址、金额、总结，判为 user_filled。",
                        "如果段落本身是占位形式、下划线、破折号、待填写提示，判为 blank_slot。",
                        "输出格式必须为 {classification, reason}。",
                    ],
                    "paragraph": paragraph_payload,
                },
                ensure_ascii=False,
            ),
        },
    ]


def _normalize_block_result(raw_result):
    if isinstance(raw_result, str):
        raw_text = raw_result.strip()
        try:
            raw_result = json.loads(raw_text)
        except Exception:
            return {"rows": [], "notes": raw_text}
    if not isinstance(raw_result, dict):
        return {"rows": [], "notes": str(raw_result)}

    normalized_rows = []
    for row_result in raw_result.get("rows", []):
        if isinstance(row_result, str):
            try:
                row_result = json.loads(row_result)
            except Exception:
                continue
        if not isinstance(row_result, dict):
            continue
        normalized_cells = []
        for cell_result in row_result.get("cells", []):
            if isinstance(cell_result, str):
                try:
                    cell_result = json.loads(cell_result)
                except Exception:
                    continue
            if not isinstance(cell_result, dict):
                continue
            classification = str(cell_result.get("classification", "")).strip()
            if classification not in ALLOWED_CLASSIFICATIONS:
                continue
            normalized_cells.append(
                {
                    "cell_index": cell_result.get("cell_index"),
                    "classification": classification,
                    "reason": str(cell_result.get("reason", "")).strip(),
                }
            )
        normalized_rows.append(
            {
                "row_index": row_result.get("row_index"),
                "cells": normalized_cells,
            }
        )
    return {"rows": normalized_rows, "notes": str(raw_result.get("notes", "")).strip()}


def _fallback_classification(text):
    return "blank_slot" if not _clean_text(text) else "template_fixed"


def _is_short_numeric_header_text(text):
    text = _clean_text(text)
    if not text:
        return False
    return bool(re.fullmatch(r"\d{1,2}", text))


def _is_value_like_text(text):
    text = _clean_text(text)
    if not text:
        return False
    return bool(re.fullmatch(r"\d+(?:\.\d+)?", text))


def _slugify_ascii(text):
    text = _clean_text(text).lower().replace("\n", " ")
    text = "".join(ch if ch.isascii() and (ch.isalnum() or ch == "_") else "_" for ch in text)
    while "__" in text:
        text = text.replace("__", "_")
    return text.strip("_")


def _next_placeholder_name(base_name, used_names, fallback_prefix):
    candidate = _slugify_ascii(base_name)
    if not candidate:
        candidate = fallback_prefix
    if candidate not in used_names:
        used_names.add(candidate)
        return candidate
    serial = 2
    while f"{candidate}_{serial}" in used_names:
        serial += 1
    candidate = f"{candidate}_{serial}"
    used_names.add(candidate)
    return candidate


def _call_block_llm(
    client,
    model,
    block_payload,
):
    response = client.chat.completions.create(
        model=model,
        messages=_build_block_messages(block_payload),
        temperature=0,
    )
    content = response.choices[0].message.content or ""
    return _normalize_block_result(content)


def _normalize_paragraph_result(raw_result):
    if isinstance(raw_result, str):
        raw_text = raw_result.strip()
        try:
            raw_result = json.loads(raw_text)
        except Exception:
            return {"classification": "template_fixed", "reason": raw_text}
    if not isinstance(raw_result, dict):
        return {"classification": "template_fixed", "reason": str(raw_result)}
    classification = str(raw_result.get("classification", "")).strip()
    if classification not in ALLOWED_CLASSIFICATIONS:
        classification = "template_fixed"
    return {
        "classification": classification,
        "reason": str(raw_result.get("reason", "")).strip(),
    }


def _call_block_llm_with_client(
    model,
    block_payload,
    *,
    base_url,
    api_key,
    timeout,
    connect_timeout,
    max_retries,
):
    client = OpenAI(
        api_key=api_key,
        base_url=base_url,
        timeout=httpx.Timeout(timeout, connect=connect_timeout),
        max_retries=max_retries,
    )
    return _call_block_llm(client, model, block_payload)


def _call_paragraph_llm_with_client(
    model,
    paragraph_payload,
    *,
    base_url,
    api_key,
    timeout,
    connect_timeout,
    max_retries,
):
    client = OpenAI(
        api_key=api_key,
        base_url=base_url,
        timeout=httpx.Timeout(timeout, connect=connect_timeout),
        max_retries=max_retries,
    )
    response = client.chat.completions.create(
        model=model,
        messages=_build_paragraph_messages(paragraph_payload),
        temperature=0,
    )
    content = response.choices[0].message.content or ""
    return _normalize_paragraph_result(content)


def _merge_block_votes(table):
    votes = defaultdict(Counter)
    reasons = defaultdict(list)
    cell_info = {}
    for row in table.get("rows", []):
        for cell in row.get("cells", []):
            key = (row["row_index"], cell["cell_index"])
            cell_info[key] = cell

    for block in table.get("block_results", []):
        for row_result in block.get("rows", []):
            row_index = row_result.get("row_index")
            for cell_result in row_result.get("cells", []):
                key = (row_index, cell_result.get("cell_index"))
                if key not in cell_info:
                    continue
                classification = cell_result.get("classification")
                if classification not in ALLOWED_CLASSIFICATIONS:
                    continue
                votes[key][classification] += 1
                reason = cell_result.get("reason", "").strip()
                if reason:
                    reasons[key].append(reason)

    merged_rows = []
    for row in table.get("rows", []):
        merged_cells = []
        for cell in row.get("cells", []):
            key = (row["row_index"], cell["cell_index"])
            if votes[key]:
                ordered = sorted(
                    votes[key].items(),
                    key=lambda item: (
                        -item[1],
                        0 if item[0] == "template_fixed" else 1,
                        item[0],
                    ),
                )
                classification = ordered[0][0]
                source = "llm"
            else:
                classification = _fallback_classification(cell.get("text", ""))
                source = "fallback"
            merged_cells.append(
                {
                    "cell_index": cell["cell_index"],
                    "text": cell.get("text", ""),
                    "span_width": cell.get("span_width", 1),
                    "span_height": cell.get("span_height", 1),
                    "classification": classification,
                    "reason": reasons[key][0] if reasons[key] else "",
                    "source": source,
                }
            )
        merged_rows.append({"row_index": row["row_index"], "cells": merged_cells})
    return _apply_structural_overrides(merged_rows)


def _apply_structural_overrides(rows):
    if len(rows) < 2:
        return rows

    for row_index in range(len(rows) - 1):
        current_row = rows[row_index]
        next_row = rows[row_index + 1]
        current_cells = current_row.get("cells", [])
        next_map = {cell["cell_index"]: cell for cell in next_row.get("cells", [])}

        header_candidates = [
            cell
            for cell in current_cells
            if _is_short_numeric_header_text(cell.get("text", ""))
        ]
        fixed_context_count = sum(
            1
            for cell in current_cells
            if cell.get("classification") == "template_fixed"
            and not _is_short_numeric_header_text(cell.get("text", ""))
            and _clean_text(cell.get("text", ""))
        )
        if len(header_candidates) < 3 or fixed_context_count < 1:
            continue

        aligned_value_count = 0
        for header_cell in header_candidates:
            next_cell = next_map.get(header_cell["cell_index"])
            if not next_cell:
                continue
            if next_cell.get("classification") in {"user_filled", "blank_slot"} and (
                _is_value_like_text(next_cell.get("text", "")) or not _clean_text(next_cell.get("text", ""))
            ):
                aligned_value_count += 1

        if aligned_value_count < 3:
            continue

        for header_cell in header_candidates:
            header_cell["classification"] = "template_fixed"
            header_cell["reason"] = "numeric_column_header_override"
            header_cell["source"] = "override"

    return rows


def extract_docx_template_llm(
    docx_path,
    model=DEFAULT_LLM_MODEL,
    base_url=DEFAULT_LLM_BASE_URL,
    api_key=DEFAULT_LLM_API_KEY,
    timeout=DEFAULT_LLM_TIMEOUT,
    connect_timeout=DEFAULT_LLM_CONNECT_TIMEOUT,
    max_retries=DEFAULT_LLM_MAX_RETRIES,
    block_size=DEFAULT_BLOCK_SIZE,
    block_step=DEFAULT_BLOCK_STEP,
    max_workers=DEFAULT_LLM_MAX_WORKERS,
    include_block_results=False,
):
    paragraphs = _parse_docx_paragraphs(docx_path)
    tables = _parse_docx_tables(docx_path)

    result_paragraphs = []
    if paragraphs:
        paragraph_worker_count = max(1, min(int(max_workers or DEFAULT_LLM_MAX_WORKERS), len(paragraphs)))
        if paragraph_worker_count == 1:
            for paragraph in paragraphs:
                llm_result = _call_paragraph_llm_with_client(
                    model,
                    paragraph,
                    base_url=base_url,
                    api_key=api_key,
                    timeout=timeout,
                    connect_timeout=connect_timeout,
                    max_retries=max_retries,
                )
                result_paragraphs.append(
                    {
                        "index": paragraph["index"],
                        "text": paragraph["text"],
                        "classification": llm_result.get("classification", "template_fixed"),
                        "reason": llm_result.get("reason", ""),
                        "source": "llm",
                    }
                )
        else:
            with ThreadPoolExecutor(max_workers=paragraph_worker_count) as executor:
                future_map = {}
                for paragraph in paragraphs:
                    future = executor.submit(
                        _call_paragraph_llm_with_client,
                        model,
                        paragraph,
                        base_url=base_url,
                        api_key=api_key,
                        timeout=timeout,
                        connect_timeout=connect_timeout,
                        max_retries=max_retries,
                    )
                    future_map[future] = paragraph
                for future in as_completed(future_map):
                    paragraph = future_map[future]
                    llm_result = future.result()
                    result_paragraphs.append(
                        {
                            "index": paragraph["index"],
                            "text": paragraph["text"],
                            "classification": llm_result.get("classification", "template_fixed"),
                            "reason": llm_result.get("reason", ""),
                            "source": "llm",
                        }
                    )
        result_paragraphs.sort(key=lambda item: item["index"])

    result_tables = []
    for table in tables:
        block_results = []
        blocks = _chunk_table_rows(table["rows"], block_size=block_size, block_step=block_step)
        if blocks:
            worker_count = max(1, min(int(max_workers or DEFAULT_LLM_MAX_WORKERS), len(blocks)))
            if worker_count == 1:
                for block in blocks:
                    payload = _build_block_payload(table["table_index"], block["block_index"], block["rows"])
                    llm_result = _call_block_llm_with_client(
                        model,
                        payload,
                        base_url=base_url,
                        api_key=api_key,
                        timeout=timeout,
                        connect_timeout=connect_timeout,
                        max_retries=max_retries,
                    )
                    block_results.append(
                        {
                            "block_index": block["block_index"],
                            "row_indices": [row["row_index"] for row in block["rows"]],
                            "rows": llm_result.get("rows", []),
                            "notes": llm_result.get("notes", ""),
                        }
                    )
            else:
                with ThreadPoolExecutor(max_workers=worker_count) as executor:
                    future_map = {}
                    for block in blocks:
                        payload = _build_block_payload(table["table_index"], block["block_index"], block["rows"])
                        future = executor.submit(
                            _call_block_llm_with_client,
                            model,
                            payload,
                            base_url=base_url,
                            api_key=api_key,
                            timeout=timeout,
                            connect_timeout=connect_timeout,
                            max_retries=max_retries,
                        )
                        future_map[future] = block
                    for future in as_completed(future_map):
                        block = future_map[future]
                        llm_result = future.result()
                        block_results.append(
                            {
                                "block_index": block["block_index"],
                                "row_indices": [row["row_index"] for row in block["rows"]],
                                "rows": llm_result.get("rows", []),
                                "notes": llm_result.get("notes", ""),
                            }
                        )
                block_results.sort(key=lambda item: item["block_index"])
        merged_rows = _merge_block_votes({"rows": table["rows"], "block_results": block_results})
        result_tables.append(
            {
                "table_index": table["table_index"],
                "rows": merged_rows,
            }
        )
        if include_block_results:
            result_tables[-1]["block_results"] = block_results

    return {
        "docx_path": docx_path,
        "paragraphs": result_paragraphs,
        "tables": result_tables,
    }


def _default_output_paths(docx_path, output_path=None):
    base, ext = os.path.splitext(docx_path)
    docx_output = output_path or f"{base}_llm_template{ext or '.docx'}"
    json_output = f"{os.path.splitext(docx_output)[0]}_fields.json"
    return docx_output, json_output


def _default_llm_result_json_path(docx_path, output_path=None):
    if output_path:
        return f"{os.path.splitext(output_path)[0]}_llm_result.json"
    base, _ext = os.path.splitext(docx_path)
    return f"{base}_llm_result.json"


def _build_paragraph_replacement_text(text, placeholder, classification):
    text = str(text or "")
    prefix_value_match = re.match(r"^(.{1,60}?[:：])(\s*)(.+)$", text)
    if prefix_value_match and classification in {"user_filled", "blank_slot"}:
        prefix = prefix_value_match.group(1)
        separator = prefix_value_match.group(2)
        return f"{prefix}{separator}{placeholder}"

    prefix_only_match = re.match(r"^(.{1,60}?[:：])(\s*)$", text)
    if prefix_only_match and classification == "blank_slot":
        prefix = prefix_only_match.group(1)
        separator = prefix_only_match.group(2)
        return f"{prefix}{separator}{placeholder}"

    blank_match = re.match(r"^(.{1,60}?)([_]{2,}|[—-]{2,}|[.．。]{2,})\s*$", text)
    if blank_match and classification == "blank_slot":
        prefix = blank_match.group(1).rstrip()
        if prefix:
            return f"{prefix} {placeholder}"

    return placeholder


def build_template_metadata_from_llm_result(llm_result):
    used_names = set()
    fields = []
    replacements = {}
    paragraph_replacements = {}
    for paragraph in llm_result.get("paragraphs", []):
        classification = paragraph.get("classification")
        if classification not in REPLACEABLE_CLASSIFICATIONS:
            continue
        paragraph_index = paragraph.get("index")
        fallback_prefix = f"p{paragraph_index}"
        name = _next_placeholder_name("", used_names, fallback_prefix)
        placeholder = "{{" + name + "}}"
        rendered_text = _build_paragraph_replacement_text(
            paragraph.get("text", ""),
            placeholder,
            classification,
        )
        paragraph_replacements[paragraph_index] = rendered_text
        fields.append(
            {
                "name": name,
                "placeholder": placeholder,
                "classification": classification,
                "paragraph_index": paragraph_index,
                "sample_value": paragraph.get("text", ""),
                "rendered_text": rendered_text,
            }
        )
    for table in llm_result.get("tables", []):
        table_index = table.get("table_index")
        for row in table.get("rows", []):
            row_index = row.get("row_index")
            for cell in row.get("cells", []):
                classification = cell.get("classification")
                if classification not in REPLACEABLE_CLASSIFICATIONS:
                    continue
                cell_index = cell.get("cell_index")
                fallback_prefix = f"t{table_index}_r{row_index}_c{cell_index}"
                name = _next_placeholder_name("", used_names, fallback_prefix)
                placeholder = "{{" + name + "}}"
                replacements[(table_index, row_index, cell_index)] = placeholder
                fields.append(
                    {
                        "name": name,
                        "placeholder": placeholder,
                        "classification": classification,
                        "table_index": table_index,
                        "row_index": row_index,
                        "cell_index": cell_index,
                        "sample_value": cell.get("text", ""),
                        # "reason": cell.get("reason", ""),
                    }
                )
    return {
        "docx_path": llm_result.get("docx_path", ""),
        "fields": fields,
    }, replacements, paragraph_replacements


def generate_docx_template_from_llm_result(docx_path, llm_result, output_path=None):
    metadata, replacements, paragraph_replacements = build_template_metadata_from_llm_result(llm_result)
    doc = Document(docx_path)

    for paragraph_index, replacement in paragraph_replacements.items():
        if 0 <= paragraph_index < len(doc.paragraphs):
            doc.paragraphs[paragraph_index].text = replacement

    for table_index, table in enumerate(doc.tables):
        for row_index, _row in enumerate(table.rows):
            physical_cells = {
                cell_index: cell
                for cell_index, cell, _span_width, _span_height in _iter_row_cells_with_spans(table, row_index)
            }
            for cell_index, cell in physical_cells.items():
                replacement = replacements.get((table_index, row_index, cell_index))
                if replacement:
                    cell.text = replacement

    docx_output, json_output = _default_output_paths(docx_path, output_path=output_path)
    doc.save(docx_output)
    with open(json_output, "w", encoding="utf-8") as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)

    return {
        "output_path": docx_output,
        "metadata_path": json_output,
        "metadata": metadata,
    }


def generate_docx_template_with_llm(
    docx_path,
    output_path=None,
    llm_result_json_path=None,
    model=DEFAULT_LLM_MODEL,
    base_url=DEFAULT_LLM_BASE_URL,
    api_key=DEFAULT_LLM_API_KEY,
    timeout=DEFAULT_LLM_TIMEOUT,
    connect_timeout=DEFAULT_LLM_CONNECT_TIMEOUT,
    max_retries=DEFAULT_LLM_MAX_RETRIES,
    block_size=DEFAULT_BLOCK_SIZE,
    block_step=DEFAULT_BLOCK_STEP,
    max_workers=DEFAULT_LLM_MAX_WORKERS,
    include_block_results=False,
):
    llm_result = extract_docx_template_llm(
        docx_path,
        model=model,
        base_url=base_url,
        api_key=api_key,
        timeout=timeout,
        connect_timeout=connect_timeout,
        max_retries=max_retries,
        block_size=block_size,
        block_step=block_step,
        max_workers=max_workers,
        include_block_results=include_block_results,
    )

    llm_result_json_path = llm_result_json_path or _default_llm_result_json_path(
        docx_path,
        output_path=output_path,
    )
    with open(llm_result_json_path, "w", encoding="utf-8") as f:
        json.dump(llm_result, f, ensure_ascii=False, indent=2)

    template_result = generate_docx_template_from_llm_result(
        docx_path,
        llm_result,
        output_path=output_path,
    )
    template_result["llm_result_path"] = llm_result_json_path
    template_result["llm_result"] = llm_result
    return template_result


if __name__ == "__main__":
    # import sys
    #
    # args = [arg for arg in sys.argv[1:] if arg]
    # include_block_results = False
    # generate_template = False
    # if "--include-block-results" in args:
    #     include_block_results = True
    #     args.remove("--include-block-results")
    # if "--generate-template" in args:
    #     generate_template = True
    #     args.remove("--generate-template")
    # if len(args) != 1:
    #     raise SystemExit(
    #         "Usage: python extract_doc_template_llm.py <template.docx> [--include-block-results] [--generate-template]"
    #     )
    # result = extract_docx_template_llm(
    #     r"D:\Python Project\Hermes\filled_template2.docx"
    # )
    # print(json.dumps(result, ensure_ascii=False, indent=2))
    # if generate_template:
    #     template_result = generate_docx_template_from_llm_result(args[0], result)
    #     print(json.dumps(template_result, ensure_ascii=False, indent=2))
    # else:
    #     print(json.dumps(result, ensure_ascii=False, indent=2))
    #
    # import json
    #
    # with open("./temp2.json", "r", encoding="utf-8") as f:
    #     llm_result = json.load(f)
    #
    # result = generate_docx_template_from_llm_result(
    #     r"D:\Python Project\Hermes\filled_template2.docx",
    #     llm_result,
    #     output_path=r"D:\Python Project\Hermes\filled_template_rendered2.docx"
    # )
    #
    # print(result["output_path"])
    # print(result["metadata_path"])

    result = generate_docx_template_with_llm(
        r"D:\Python Project\Hermes\filled_template2.docx",
        output_path=r"D:\Python Project\Hermes\filled_template_rendered2.docx"
    )

    print(result["llm_result_path"])
    print(result["output_path"])
    print(result["metadata_path"])
