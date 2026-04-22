import json
import os
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from statistics import median

from docx import Document
from openai import OpenAI
from docx.oxml.ns import qn

DATE_PATTERN = re.compile(r"\b\d{4}[-/.年]\d{1,2}[-/.月]\d{1,2}(?:日)?\b")
NUMBER_PATTERN = re.compile(r"\d")
MONEY_PATTERN = re.compile(r"(元|万元|亿元|人民币|￥|¥)")
PERCENT_PATTERN = re.compile(r"%|％")
ID_LIKE_PATTERN = re.compile(r"[0-9Xx]{8,}")
STRUCTURED_VALUE_PATTERN = re.compile(r"[A-Za-z0-9].*[A-Za-z0-9]")
ADDRESS_HINT_PATTERN = re.compile(r"(省|市|区|县|路|街|号|室)")

DEFAULT_LLM_MODEL = os.getenv("DOC_TEMPLATE_LLM_MODEL", "qwen-max")
DEFAULT_LLM_BASE_URL = os.getenv("DOC_TEMPLATE_LLM_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1").strip()
DEFAULT_LLM_API_KEY = os.getenv("DOC_TEMPLATE_LLM_API_KEY", "sk-0b4c79748dc04bd8bebbadde4a566fba").strip()
DEFAULT_LLM_MAX_WORKERS = max(1, int(os.getenv("DOC_TEMPLATE_LLM_MAX_WORKERS", "8")))
DEFAULT_LLM_BATCH_SIZE = max(1, int(os.getenv("DOC_TEMPLATE_LLM_BATCH_SIZE", "8")))


def _clean_text(text):
    return re.sub(r"\s+", " ", (text or "")).strip()


def _unique_preserve_order(items):
    seen = set()
    result = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        result.append(item)
    return result


def _dedupe_field_pairs(field_pairs):
    """Drop duplicate pairs produced by vertically merged cells."""
    seen = set()
    result = []
    for pair in field_pairs:
        key = (
            pair.get("table_index"),
            pair.get("label_cell_index"),
            pair.get("value_cell_index"),
            pair.get("label"),
            pair.get("value"),
        )
        if key in seen:
            continue
        seen.add(key)
        result.append(pair)
    return result


def _is_numeric_like_text(text):
    text = text.strip()
    return bool(text) and bool(re.fullmatch(r"[\d.\-/%％年月日天个万元亿元XxA-Za-z]+", text))


def _iter_unique_row_cells(row):
    """Yield only unique physical cells for a row.

    python-docx expands merged cells across the row grid, so repeated entries in
    row.cells can point to the same underlying XML cell. We keep only the first
    occurrence of each physical cell to avoid duplicate output.
    """
    seen_tc_ids = set()
    for grid_index, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id in seen_tc_ids:
            continue
        seen_tc_ids.add(tc_id)
        yield grid_index, cell


def _get_vmerge_state(cell):
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    vmerge = tc_pr.find(qn("w:vMerge"))
    if vmerge is None:
        return None
    return vmerge.get(qn("w:val")) or "continue"


def _compute_span_height(table, row_index, grid_index, span_width):
    """Estimate vertical span height for a cell starting at the given grid position."""
    current_row = table.rows[row_index]
    current_cell = current_row.cells[grid_index]
    current_vmerge = _get_vmerge_state(current_cell)
    if current_vmerge not in {"restart", "continue"}:
        return 1

    height = 1
    for next_row_index in range(row_index + 1, len(table.rows)):
        next_row = table.rows[next_row_index]
        if grid_index >= len(next_row.cells):
            break
        next_cell = next_row.cells[grid_index]
        next_vmerge = _get_vmerge_state(next_cell)
        if next_vmerge not in {"restart", "continue"}:
            break
        height += 1
        if span_width > 1:
            end_index = min(grid_index + span_width, len(next_row.cells))
            if not all(_get_vmerge_state(next_row.cells[i]) in {"restart", "continue"} for i in range(grid_index, end_index)):
                break
    return height


def _iter_row_cells_with_spans(table, row_index):
    """Yield unique physical cells with horizontal and vertical span metadata."""
    row = table.rows[row_index]
    raw_cells = list(row.cells)
    seen_tc_ids = set()
    grid_index = 0

    while grid_index < len(raw_cells):
        cell = raw_cells[grid_index]
        tc_id = id(cell._tc)
        if tc_id in seen_tc_ids:
            grid_index += 1
            continue

        span_width = 1
        next_index = grid_index + 1
        while next_index < len(raw_cells) and id(raw_cells[next_index]._tc) == tc_id:
            span_width += 1
            next_index += 1

        seen_tc_ids.add(tc_id)
        span_height = _compute_span_height(table, row_index, grid_index, span_width)
        yield grid_index, cell, span_width, span_height
        grid_index = next_index


def _score_cell_text(text):
    """Return fixed/variable scores and reasons for a single cell."""
    fixed_score = 0
    variable_score = 0
    reasons = []
    length = len(text)

    if not text:
        variable_score += 3
        reasons.append("empty_cell")
        return fixed_score, variable_score, reasons

    if text.endswith((":", "：")):
        fixed_score += 3
        reasons.append("label_suffix")

    if length <= 8:
        fixed_score += 2
        reasons.append("short_text")
    elif length <= 16:
        fixed_score += 1
        reasons.append("medium_short_text")
    else:
        variable_score += 1
        reasons.append("long_text")

    if NUMBER_PATTERN.search(text):
        variable_score += 2
        reasons.append("contains_digits")

    if DATE_PATTERN.search(text):
        variable_score += 3
        reasons.append("date_like")

    if MONEY_PATTERN.search(text):
        variable_score += 2
        reasons.append("money_like")

    if PERCENT_PATTERN.search(text):
        variable_score += 2
        reasons.append("percent_like")

    if ID_LIKE_PATTERN.search(text):
        variable_score += 2
        reasons.append("id_like")

    if ADDRESS_HINT_PATTERN.search(text) and length >= 10:
        variable_score += 2
        reasons.append("address_like")

    if any(mark in text for mark in ("，", ",", "；", ";", "、")) and length >= 10:
        variable_score += 1
        reasons.append("descriptive_value_like")

    if STRUCTURED_VALUE_PATTERN.search(text) and NUMBER_PATTERN.search(text):
        variable_score += 1
        reasons.append("structured_value_like")

    if length <= 12 and not NUMBER_PATTERN.search(text) and not any(mark in text for mark in ("，", ",", "；", ";")):
        fixed_score += 1
        reasons.append("compact_label_like")

    return fixed_score, variable_score, reasons


def _classify_cell(text, row_lengths, non_empty_lengths, span_width=1, span_height=1):
    """Classify a cell as fixed/variable/unknown using text and row-relative signals."""
    fixed_score, variable_score, reasons = _score_cell_text(text)
    length = len(text)

    if span_height > 1:
        fixed_score += 2
        reasons.append("row_spanning_cell")

    if span_width > 1 and text:
        if length <= 20:
            fixed_score += 1
            reasons.append("short_wide_cell")
        else:
            variable_score += 1
            reasons.append("wide_value_area")

    if span_width > 1 and not text:
        variable_score += 3
        reasons.append("empty_wide_cell")

    if non_empty_lengths:
        min_len = min(non_empty_lengths)
        max_len = max(non_empty_lengths)
        med_len = median(non_empty_lengths)

        if text:
            if length == min_len and max_len >= max(10, length + 4):
                fixed_score += 2
                reasons.append("shortest_in_row")
            if length == max_len and length >= min_len + 4:
                variable_score += 2
                reasons.append("longest_in_row")
            if length < med_len:
                fixed_score += 1
                reasons.append("shorter_than_row_median")
            elif length > med_len:
                variable_score += 1
                reasons.append("longer_than_row_median")

    classification = "unknown"
    confidence = abs(variable_score - fixed_score)
    if variable_score >= fixed_score + 2:
        classification = "variable"
    elif fixed_score >= variable_score + 2:
        classification = "fixed"

    return {
        "classification": classification,
        "score_fixed": fixed_score,
        "score_variable": variable_score,
        "confidence": confidence,
        "reasons": _unique_preserve_order(reasons),
    }


def _build_cell_payload(
    table_index,
    row_index,
    cell_index,
    text,
    row_lengths,
    non_empty_lengths,
    span_width=1,
    span_height=1,
):
    classification = _classify_cell(text, row_lengths, non_empty_lengths, span_width=span_width, span_height=span_height)
    return {
        "type": "table_cell",
        "raw_text": text,
        "table_index": table_index,
        "row_index": row_index,
        "cell_index": cell_index,
        "span_width": span_width,
        "span_height": span_height,
        **classification,
    }


def _coerce_cell_classification(value):
    value = str(value or "").strip().lower()
    if value in {"fixed", "variable", "unknown"}:
        return value
    return "unknown"


def _is_row_heading(cells):
    """Treat a row as a heading when everything looks fixed and no cell looks variable."""
    meaningful = [
        cell
        for cell in cells
        if cell["raw_text"] or (cell["classification"] == "variable" and cell.get("span_width", 1) > 1)
    ]
    non_empty = [cell for cell in meaningful if cell["raw_text"]]
    if not meaningful:
        return False
    if all(cell["classification"] in {"fixed", "unknown"} for cell in meaningful) and not any(
        cell["classification"] == "variable" for cell in meaningful
    ):
        max_len = max(len(cell["raw_text"]) for cell in non_empty)
        return len(non_empty) == 1 or (len(non_empty) <= 2 and max_len <= 20)
    return False


def _extract_row_field_pairs(cells):
    """Extract likely fixed-variable pairs from a row."""
    def _is_group_heading_cell(cell):
        text = cell["raw_text"].strip()
        return (
            bool(text)
            and cell.get("span_height", 1) > 1
            and len(text) <= 12
            and not any(mark in text for mark in ("，", ",", "；", ";", "、"))
        )

    raw_non_empty = [
        cell
        for cell in cells
        if (cell["raw_text"] or (cell["classification"] == "variable" and cell.get("span_width", 1) > 1))
        and not _is_group_heading_cell(cell)
    ]
    if len(raw_non_empty) < 2 or _is_row_heading(cells):
        return []

    # In multi-field rows, the first vertically merged cell is usually a section
    # group label rather than a true field label. Keep it out of pairing to
    # avoid matches like "专精特新... -> 从事特定细分市场年限".
    non_empty = list(raw_non_empty)
    if (
        len(non_empty) >= 3
        and non_empty[0].get("span_height", 1) > 1
        and non_empty[0].get("cell_index") == 0
    ):
        non_empty = non_empty[1:]
    if len(non_empty) < 2:
        return []

    def _is_value_like(cell):
        text = cell["raw_text"].strip()
        if not text:
            return cell["classification"] == "variable" and cell.get("span_width", 1) > 1
        if cell["classification"] == "variable":
            return True
        if text in {"是", "否", "有", "无"}:
            return True
        if DATE_PATTERN.search(text) or MONEY_PATTERN.search(text) or PERCENT_PATTERN.search(text):
            return True
        if ID_LIKE_PATTERN.search(text) or ADDRESS_HINT_PATTERN.search(text):
            return True
        if _is_numeric_like_text(text):
            return True
        if any(mark in text for mark in ("，", ",", "；", ";", "、")) and len(text) >= 6:
            return True
        if len(text) >= 10 and any(mark in text for mark in ("，", ",", "；", ";", "、")):
            return True
        if len(text) >= 12 and cell.get("span_width", 1) > 1:
            return True
        # Medium-length Chinese phrases in a wide cell are often filled values
        # such as company names, stage names, business types, etc.
        if (
            cell.get("span_width", 1) > 1
            and 4 <= len(text) <= 24
            and re.fullmatch(r"[\u4e00-\u9fffA-Za-z0-9（）()·\-]+", text)
            and not any(marker in text for marker in ("是否", "意见", "说明", "情况"))
        ):
            return True
        return False

    def _is_label_like(cell):
        text = cell["raw_text"].strip()
        if not text:
            return False
        if _is_group_heading_cell(cell):
            return False
        if any(marker in text for marker in ("是否", "是否达", "是否高于", "是否低于", "是否超过", "是否存在")):
            return True
        # Long Chinese text in a narrow single-column cell is often still a
        # field label, especially for stage/type/basis/style descriptions.
        if (
            cell.get("span_width", 1) == 1
            and 8 <= len(text) <= 32
            and not NUMBER_PATTERN.search(text)
            and not any(mark in text for mark in ("，", ",", "；", ";", "、"))
            and re.fullmatch(r"[\u4e00-\u9fffA-Za-z（）()·\-]+", text)
        ):
            return True
        if _is_numeric_like_text(text) and text not in {"是否"}:
            return False
        if cell["classification"] == "fixed":
            return True
        if cell["classification"] == "unknown":
            if len(text) <= 20 and not any(mark in text for mark in ("，", ",", "；", ";", "、")):
                return True
        return False

    pairs = []
    used_indices = set()

    # Single label + single large value/description row.
    if len(non_empty) == 2:
        left, right = non_empty
        if _is_label_like(left) and _is_value_like(right):
            return [
                {
                    "label": left["raw_text"],
                    "label_cell_index": left["cell_index"],
                    "value": right["raw_text"],
                    "value_cell_index": right["cell_index"],
                }
            ]

    # Alternating label/value rows: label value label value ...
    alternating_pairs = []
    alternating_ok = len(non_empty) >= 4
    for idx, cell in enumerate(non_empty):
        if idx % 2 == 0:
            if not _is_label_like(cell):
                alternating_ok = False
                break
        else:
            if not _is_value_like(cell):
                alternating_ok = False
                break
    if alternating_ok:
        for idx in range(0, len(non_empty) - 1, 2):
            left = non_empty[idx]
            right = non_empty[idx + 1]
            alternating_pairs.append(
                {
                    "label": left["raw_text"],
                    "label_cell_index": left["cell_index"],
                    "value": right["raw_text"],
                    "value_cell_index": right["cell_index"],
                }
            )
        if alternating_pairs:
            return alternating_pairs

    # Prefer adjacent label-value pairs inside the row. This captures repeated
    # patterns like label/value/label/value more reliably than a single global
    # pending_label state.
    for idx in range(len(non_empty) - 1):
        left = non_empty[idx]
        right = non_empty[idx + 1]
        if idx in used_indices or (idx + 1) in used_indices:
            continue
        if _is_label_like(left) and (
            _is_value_like(right)
            or (
                len(left["raw_text"].strip()) >= 10
                and len(right["raw_text"].strip()) <= 12
                and right.get("span_width", 1) > 1
                and re.fullmatch(r"[\u4e00-\u9fffA-Za-z0-9（）()·\-]+", right["raw_text"].strip() or "")
                and not any(marker in right["raw_text"] for marker in ("是否", "意见", "说明", "情况"))
            )
        ):
            pairs.append(
                {
                    "label": left["raw_text"],
                    "label_cell_index": left["cell_index"],
                    "value": right["raw_text"],
                    "value_cell_index": right["cell_index"],
                }
            )
            used_indices.add(idx)
            used_indices.add(idx + 1)

    # Fallback: keep the old state-machine idea, but only on cells not already
    # paired and with stricter label/value tests.
    pending_label = None
    for idx, cell in enumerate(non_empty):
        if idx in used_indices:
            continue
        if _is_label_like(cell):
            pending_label = cell
            continue
        if pending_label is not None and _is_value_like(cell):
            pairs.append(
                {
                    "label": pending_label["raw_text"],
                    "label_cell_index": pending_label["cell_index"],
                    "value": cell["raw_text"],
                    "value_cell_index": cell["cell_index"],
                }
            )
            pending_label = None

    return pairs


def _detect_matrix_block(header_row, value_row):
    """Detect a two-row horizontal matrix like month headers + numeric values."""
    header_cells = [cell for cell in header_row["cells"] if cell["raw_text"]]
    value_cells = [cell for cell in value_row["cells"] if cell["raw_text"]]
    if len(header_cells) < 3 or len(value_cells) < 3:
        return None

    group_label_cell = None
    header_candidates = header_cells
    value_candidates = value_cells

    spanning_labels = [
        cell for cell in header_cells
        if cell.get("span_height", 1) > 1 and not _is_numeric_like_text(cell["raw_text"])
    ]
    if len(spanning_labels) >= 2:
        group_label_cell = spanning_labels[1]
        header_candidates = [cell for cell in header_cells if cell["cell_index"] > group_label_cell["cell_index"]]
        value_candidates = [cell for cell in value_cells if cell["cell_index"] > group_label_cell["cell_index"]]
    elif spanning_labels:
        group_label_cell = spanning_labels[0]
        header_candidates = [cell for cell in header_cells if cell["cell_index"] > group_label_cell["cell_index"]]
        value_candidates = [cell for cell in value_cells if cell["cell_index"] > group_label_cell["cell_index"]]
    elif not _is_numeric_like_text(header_cells[0]["raw_text"]):
        group_label_cell = header_cells[0]
        header_candidates = header_cells[1:]
        value_candidates = value_cells[1:]

    if not group_label_cell or len(header_candidates) < 3 or len(value_candidates) < 3:
        return None

    aligned_headers = []
    aligned_values = []
    value_map = {cell["cell_index"]: cell for cell in value_candidates}
    for cell in header_candidates:
        if cell["cell_index"] not in value_map:
            continue
        aligned_headers.append(cell)
        aligned_values.append(value_map[cell["cell_index"]])

    if len(aligned_headers) < 3:
        return None

    if not all(_is_numeric_like_text(cell["raw_text"]) for cell in aligned_headers):
        return None
    if not all(_is_numeric_like_text(cell["raw_text"]) for cell in aligned_values):
        return None

    entries = []
    expanded_pairs = []
    label_text = group_label_cell["raw_text"]
    for header_cell, value_cell in zip(aligned_headers, aligned_values):
        header_text = header_cell["raw_text"]
        value_text = value_cell["raw_text"]
        entries.append(
            {
                "header": header_text,
                "header_cell_index": header_cell["cell_index"],
                "value": value_text,
                "value_cell_index": value_cell["cell_index"],
            }
        )
        expanded_pairs.append(
            {
                "label": f"{label_text}-{header_text}",
                "label_cell_index": group_label_cell["cell_index"],
                "value": value_text,
                "value_cell_index": value_cell["cell_index"],
            }
        )

    return {
        "type": "matrix_block",
        "group_label": label_text,
        "group_label_cell_index": group_label_cell["cell_index"],
        "header_row_index": header_row["row_index"],
        "value_row_index": value_row["row_index"],
        "headers": [entry["header"] for entry in entries],
        "values": [entry["value"] for entry in entries],
        "entries": entries,
        "expanded_pairs": expanded_pairs,
    }


def _extract_table_matrix_blocks(rows):
    matrix_blocks = []
    used_row_indices = set()
    for idx in range(len(rows) - 1):
        header_row = rows[idx]
        value_row = rows[idx + 1]
        if header_row["row_index"] in used_row_indices or value_row["row_index"] in used_row_indices:
            continue
        block = _detect_matrix_block(header_row, value_row)
        if not block:
            continue
        matrix_blocks.append(block)
        used_row_indices.add(header_row["row_index"])
        used_row_indices.add(value_row["row_index"])
        header_row["row_type"] = "matrix_header_row"
        value_row["row_type"] = "matrix_value_row"
        header_row["field_pairs"] = []
        value_row["field_pairs"] = []
    return matrix_blocks, used_row_indices


def _build_blocks_for_table(table_index, rows, matrix_blocks):
    """Build semantic blocks from parsed rows for downstream extraction/judging."""
    matrix_row_indices = set()
    blocks = []

    for block_index, matrix_block in enumerate(matrix_blocks):
        matrix_row_indices.add(matrix_block["header_row_index"])
        matrix_row_indices.add(matrix_block["value_row_index"])
        blocks.append(
            {
                "block_id": f"t{table_index}_matrix_{block_index}",
                "table_index": table_index,
                "row_range": [matrix_block["header_row_index"], matrix_block["value_row_index"]],
                "block_type": "matrix_block",
                "cells": [
                    cell
                    for row in rows
                    if row["row_index"] in {matrix_block["header_row_index"], matrix_block["value_row_index"]}
                    for cell in row["cells"]
                ],
                "candidate_pairs": list(matrix_block.get("expanded_pairs", [])),
                "matrix_block": matrix_block,
            }
        )

    current_cells = []
    current_rows = []
    current_candidates = []
    current_type = None
    serial = 0

    def flush_current_block():
        nonlocal current_cells, current_rows, current_candidates, current_type, serial
        if not current_rows:
            return
        blocks.append(
            {
                "block_id": f"t{table_index}_{current_type}_{serial}",
                "table_index": table_index,
                "row_range": [current_rows[0], current_rows[-1]],
                "block_type": current_type,
                "cells": list(current_cells),
                "candidate_pairs": _dedupe_field_pairs(list(current_candidates)),
            }
        )
        serial += 1
        current_cells = []
        current_rows = []
        current_candidates = []
        current_type = None

    for row in rows:
        row_index = row["row_index"]
        if row_index in matrix_row_indices:
            flush_current_block()
            continue

        row_cells = [cell for cell in row["cells"] if cell["raw_text"]]
        non_empty_count = len(row_cells)
        has_empty_variable = any(
            (not cell["raw_text"]) and cell["classification"] == "variable"
            for cell in row["cells"]
        )
        has_pairs = bool(row["field_pairs"])

        row_block_type = "field_block"
        if row["row_type"] == "heading_row":
            row_block_type = "heading_block"
        elif has_empty_variable:
            row_block_type = "empty_slot_block"
        elif non_empty_count == 2 and any(len(cell["raw_text"]) >= 16 for cell in row_cells):
            row_block_type = "description_block"
        elif not has_pairs and non_empty_count <= 2:
            row_block_type = "description_block"

        if current_type is None:
            current_type = row_block_type
        elif current_type != row_block_type:
            flush_current_block()
            current_type = row_block_type

        current_rows.append(row_index)
        current_cells.extend(row["cells"])
        current_candidates.extend(row["field_pairs"])

    flush_current_block()
    return blocks


def _build_document_blocks(tables):
    """Build document-level blocks from parsed tables."""
    blocks = []
    for table in tables:
        blocks.extend(_build_blocks_for_table(table["table_index"], table["rows"], table.get("matrix_blocks", [])))
    return blocks


def _build_llm_row_payload(table_index, row_index, cells):
    return {
        "table_index": table_index,
        "row_index": row_index,
        "cells": [
            {
                "cell_index": cell["cell_index"],
                "text": cell["raw_text"],
                "span_width": cell["span_width"],
                "span_height": cell["span_height"],
                "heuristic_classification": cell["classification"],
                "score_fixed": cell["score_fixed"],
                "score_variable": cell["score_variable"],
                "reasons": cell["reasons"],
            }
            for cell in cells
        ],
    }


def _build_llm_messages(row_payload):
    system_prompt = (
        "你是一个文档表格解析助手。"
        "任务是把标准表格中的单元格分类为 fixed 或 variable。"
        "fixed 表示固定文本、标题、标签、提示语、说明项。"
        "variable 表示被填写的数据值、待填写内容、空白值区域。"
        "unknown 只在确实无法判断时使用。"
        "优先利用同一行内的相对关系判断，不要依赖特定字段词典。"
        "输出必须是 JSON，不要包含 markdown。"
    )
    user_prompt = {
        "task": "classify_table_row",
        "instructions": [
            "阅读这一行表格单元格及其启发式结果。",
            "返回每个 cell_index 的最终 classification。",
            "如果存在明显的标签和值关系，提取 field_pairs。",
            "row_type 只能是 heading_row、field_row、description_row 之一。",
            "仅返回 JSON 对象，格式为 {row_type, cells, field_pairs, notes}。",
        ],
        "row": row_payload,
    }
    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": json.dumps(user_prompt, ensure_ascii=False)},
    ]


def _call_row_llm(client, model, row_payload):
    response = client.chat.completions.create(
        model=model,
        messages=_build_llm_messages(row_payload),
        temperature=0,
        response_format={"type": "json_object"},
    )
    content = response.choices[0].message.content or "{}"
    return json.loads(content)


def _call_row_llm_with_client(model, base_url, api_key, row_payload):
    client = OpenAI(api_key=api_key, base_url=base_url)
    return _call_row_llm(client, model, row_payload)


def _chunked(items, chunk_size):
    for i in range(0, len(items), chunk_size):
        yield items[i:i + chunk_size]


def _apply_llm_row_result(row, llm_result):
    cell_map = {cell["cell_index"]: cell for cell in row["cells"]}
    for cell_result in llm_result.get("cells", []):
        cell_index = cell_result.get("cell_index")
        if cell_index not in cell_map:
            continue
        cell = cell_map[cell_index]
        cell["classification"] = _coerce_cell_classification(cell_result.get("classification"))
        cell["llm_reason"] = str(cell_result.get("reason", "")).strip()
        cell["classification_source"] = "llm"

    row_type = str(llm_result.get("row_type", "")).strip()
    if row_type in {"heading_row", "field_row", "description_row"}:
        row["row_type"] = row_type

    llm_pairs = []
    for pair in llm_result.get("field_pairs", []):
        label_index = pair.get("label_cell_index")
        value_index = pair.get("value_cell_index")
        if label_index not in cell_map or value_index not in cell_map:
            continue
        llm_pairs.append(
            {
                "label": cell_map[label_index]["raw_text"],
                "label_cell_index": label_index,
                "value": cell_map[value_index]["raw_text"],
                "value_cell_index": value_index,
            }
        )
    if llm_pairs:
        row["field_pairs"] = llm_pairs

    row["llm_notes"] = str(llm_result.get("notes", "")).strip()
    return row


def _refine_rows_with_llm(
    tables,
    model=DEFAULT_LLM_MODEL,
    base_url=DEFAULT_LLM_BASE_URL,
    api_key=DEFAULT_LLM_API_KEY,
    max_workers=DEFAULT_LLM_MAX_WORKERS,
    batch_size=DEFAULT_LLM_BATCH_SIZE,
):
    if not base_url or not api_key:
        raise ValueError("LLM refinement requires DOC_TEMPLATE_LLM_BASE_URL and DOC_TEMPLATE_LLM_API_KEY.")

    row_jobs = []
    for table in tables:
        for row in table["rows"]:
            row_jobs.append(
                {
                    "table_index": table["table_index"],
                    "row_index": row["row_index"],
                    "payload": _build_llm_row_payload(table["table_index"], row["row_index"], row["cells"]),
                }
            )

    if not row_jobs:
        return tables

    for batch in _chunked(row_jobs, batch_size):
        with ThreadPoolExecutor(max_workers=min(max_workers, len(batch))) as executor:
            future_map = {
                executor.submit(
                    _call_row_llm_with_client,
                    model,
                    base_url,
                    api_key,
                    job["payload"],
                ): job
                for job in batch
            }
            for future in as_completed(future_map):
                job = future_map[future]
                row = tables[job["table_index"]]["rows"][job["row_index"]]
                try:
                    llm_result = future.result()
                    _apply_llm_row_result(row, llm_result)
                except Exception as exc:
                    row["llm_error"] = str(exc)
                    row["llm_notes"] = "llm_refinement_failed_fallback_to_heuristic"

    for table in tables:
        table_pairs = []
        for row in table["rows"]:
            for pair in row["field_pairs"]:
                pair["table_index"] = table["table_index"]
                pair["row_index"] = row["row_index"]
            table_pairs.extend(row["field_pairs"])
        table["field_pairs"] = table_pairs
    return tables


def _classify_paragraph(index, text):
    """Paragraphs in this workflow are usually headings/instructions, so default to fixed."""
    text = _clean_text(text)
    fixed_score, variable_score, reasons = _score_cell_text(text)
    fixed_score += 1
    reasons.append("paragraph_defaults_to_fixed")
    classification = "fixed" if fixed_score >= variable_score else "unknown"
    return {
        "type": "paragraph",
        "index": index,
        "raw_text": text,
        "classification": classification,
        "score_fixed": fixed_score,
        "score_variable": variable_score,
        "confidence": abs(variable_score - fixed_score),
        "reasons": _unique_preserve_order(reasons),
    }


def parse_docx_to_ir(docx_path):
    """Parse a DOCX file into an intermediate representation (IR)."""
    doc = Document(docx_path)

    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = _clean_text(paragraph.text)
        if not text:
            continue
        paragraphs.append(_classify_paragraph(index, text))

    tables = []
    all_field_pairs = []
    all_matrix_blocks = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        table_field_pairs = []

        for row_index, row in enumerate(table.rows):
            raw_cells = []
            for cell_index, cell, span_width, span_height in _iter_row_cells_with_spans(table, row_index):
                text = _clean_text("\n".join(paragraph.text for paragraph in cell.paragraphs))
                raw_cells.append((cell_index, text, span_width, span_height))

            row_lengths = [len(text) for _, text, _, _ in raw_cells]
            non_empty_lengths = [length for length in row_lengths if length > 0]

            cells = [
                _build_cell_payload(
                    table_index=table_index,
                    row_index=row_index,
                    cell_index=cell_index,
                    text=text,
                    row_lengths=row_lengths,
                    non_empty_lengths=non_empty_lengths,
                    span_width=span_width,
                    span_height=span_height,
                )
                for cell_index, text, span_width, span_height in raw_cells
            ]

            row_type = "field_row"
            if _is_row_heading(cells):
                row_type = "heading_row"

            field_pairs = _extract_row_field_pairs(cells)
            for pair in field_pairs:
                pair.update({"table_index": table_index, "row_index": row_index})
            table_field_pairs.extend(field_pairs)

            rows.append(
                {
                    "row_index": row_index,
                    "row_type": row_type,
                    "cells": cells,
                    "field_pairs": field_pairs,
                }
            )

        matrix_blocks, matrix_row_indices = _extract_table_matrix_blocks(rows)
        for block in matrix_blocks:
            block["table_index"] = table_index
        all_matrix_blocks.extend(matrix_blocks)

        table_field_pairs = [
            pair
            for pair in table_field_pairs
            if pair["row_index"] not in matrix_row_indices
        ]
        table_field_pairs = _dedupe_field_pairs(table_field_pairs)
        tables.append(
            {
                "type": "table",
                "table_index": table_index,
                "rows": rows,
                "field_pairs": table_field_pairs,
                "matrix_blocks": matrix_blocks,
            }
        )
        all_field_pairs.extend(table_field_pairs)

    blocks = _build_document_blocks(tables)
    return {
        "paragraphs": paragraphs,
        "tables": tables,
        "blocks": blocks,
        "field_pairs": _dedupe_field_pairs(all_field_pairs),
        "matrix_blocks": all_matrix_blocks,
    }


def extract_from_ir(ir, use_llm=False, llm_model=DEFAULT_LLM_MODEL, llm_base_url=DEFAULT_LLM_BASE_URL, llm_api_key=DEFAULT_LLM_API_KEY, llm_max_workers=DEFAULT_LLM_MAX_WORKERS, llm_batch_size=DEFAULT_LLM_BATCH_SIZE):
    """Extract final results from IR, optionally refining rows with LLM."""
    tables = ir["tables"]
    if use_llm:
        tables = _refine_rows_with_llm(
            tables,
            model=llm_model,
            base_url=llm_base_url,
            api_key=llm_api_key,
            max_workers=llm_max_workers,
            batch_size=llm_batch_size,
        )

    all_field_pairs = []
    all_matrix_blocks = []
    for table in tables:
        all_field_pairs.extend(table["field_pairs"])
        all_matrix_blocks.extend(table.get("matrix_blocks", []))

    return {
        "tables": tables,
        "blocks": _build_document_blocks(tables),
        "field_pairs": _dedupe_field_pairs(all_field_pairs),
        "matrix_blocks": all_matrix_blocks,
    }


def extract_docx_template(
    docx_path,
    use_llm=False,
    llm_model=DEFAULT_LLM_MODEL,
    llm_base_url=DEFAULT_LLM_BASE_URL,
    llm_api_key=DEFAULT_LLM_API_KEY,
    llm_max_workers=DEFAULT_LLM_MAX_WORKERS,
    llm_batch_size=DEFAULT_LLM_BATCH_SIZE,
):
    """Parse a DOCX file and extract table semantics."""
    ir = parse_docx_to_ir(docx_path)
    extracted = extract_from_ir(
        ir,
        use_llm=use_llm,
        llm_model=llm_model,
        llm_base_url=llm_base_url,
        llm_api_key=llm_api_key,
        llm_max_workers=llm_max_workers,
        llm_batch_size=llm_batch_size,
    )
    return {
        "paragraphs": ir["paragraphs"],
        "tables": extracted["tables"],
        "blocks": extracted["blocks"],
        "field_pairs": extracted["field_pairs"],
        "matrix_blocks": extracted["matrix_blocks"],
    }


if __name__ == "__main__":
    # import sys
    #
    # args = [arg for arg in sys.argv[1:] if arg]
    # use_llm = False
    # if "--use-llm" in args:
    #     use_llm = True
    #     args.remove("--use-llm")
    #
    # if len(args) != 1:
    #     raise SystemExit("Usage: python extract_doc_template.py <template.docx> [--use-llm]")

    # result = extract_docx_template(args[0], use_llm=use_llm)
    result = extract_docx_template("D:\\Python Project\\Hermes\\filled-专精特新类公司调查报告.docx")
    print(json.dumps(result, ensure_ascii=False, indent=2))
